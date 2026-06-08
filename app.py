import streamlit as st
import os
import json
import time
import re
import base64
from io import BytesIO
from PIL import Image, ImageEnhance
import pypdfium2 as pdfium
from docx import Document
import pandas as pd
from groq import Groq

# ==============================================================================
# 🎨 BRAND THEME LAYOUT (CSS Injection)
# ==============================================================================
extractor_theme_css = """
<style>
.stApp { background-color: #FAFAFA; color: #2F3E46; }
h1, h2, h3, .stSubheader { color: #005F73 !important; font-family: 'Helvetica Neue', Arial, sans-serif; font-weight: 600; }
div.stButton > button:first-child { background-color: #005F73 !important; color: white !important; border-radius: 4px; border: none; font-weight: bold; padding: 0.5rem 2rem; transition: all 0.3s ease; }
div.stButton > button:first-child:hover { background-color: #0A9396 !important; box-shadow: 0px 4px 10px rgba(0, 95, 115, 0.2); }
section[data-testid="stSidebar"] { background-color: #E0F2F1 !important; border-right: 1px solid #B2DFDB; }
div[data-testid="stMetricValue"] { color: #005F73 !important; }
.status-box { padding: 15px; border-radius: 5px; background-color: white; border-left: 5px solid #005F73; margin-bottom: 10px; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
</style>
"""

st.set_page_config(page_title="Text Extractor Hub", page_icon="🔐", layout="wide")
st.markdown(extractor_theme_css, unsafe_allow_html=True)

# Safely check for API key (Fixes the disappearing upload button)
api_key_available = "GROQ_API_KEY" in st.secrets

# ==============================================================================
# ⚙️ SIDEBAR CONFIGURATION
# ==============================================================================
st.sidebar.markdown("<h2 style='margin-top:0;'>⚙️ Section Configurator</h2>", unsafe_allow_html=True)
st.sidebar.markdown("Define your sections and parameters below using clean JSON format:")

default_blueprint = {
    "FOR OFFICE USE": ["Receipt Date"],
    "PART 1 - PROPOSER": [
        "Name", "Date of Birth", "Aadhaar Card No", 
        "PAN Card / Unique Identification No", 
        "Proposer Permanent Residential Address including Pincode"
    ]
}

blueprint_text = st.sidebar.text_area(
    "Modify Sections & Fields:", 
    value=json.dumps(default_blueprint, indent=4), 
    height=250
)

try:
    SECTION_STRUCTURE = json.loads(blueprint_text)
    if not isinstance(SECTION_STRUCTURE, dict):
        SECTION_STRUCTURE = default_blueprint
except Exception:
    SECTION_STRUCTURE = default_blueprint

MODEL_SELECTION = st.sidebar.selectbox("Choose AI Vision Engine:", ["Llama 4 Scout Vision"])
MODEL_NAME = "meta-llama/llama-4-scout-17b-16e-instruct"

ALL_FLAT_COLUMNS = ["File Name", "Validation Report"]
for section, items in SECTION_STRUCTURE.items():
    if isinstance(items, list):
        ALL_FLAT_COLUMNS.extend(items)

st.markdown("<h1>🔐 Document Text Extractor Intelligence Hub</h1>", unsafe_allow_html=True)
st.caption("Securely parse unstructured forms into custom databases using Llama 4 Vision on Groq.")

if not api_key_available:
    st.warning("⚠️ Missing GROQ_API_KEY. Please add it to your Streamlit Advanced Settings -> Secrets to unlock processing.")

# ==============================================================================
# 🛠️ PROCESSING PIPELINE ENGINE
# ==============================================================================
def enhance_image_for_ai(img):
    """Boosts contrast and heavily increases sharpness in full color."""
    img = img.convert('RGB')
    
    # Boost Contrast
    enhancer_contrast = ImageEnhance.Contrast(img)
    img = enhancer_contrast.enhance(2.0)
    
    # Heavily Increase Sharpness
    enhancer_sharpness = ImageEnhance.Sharpness(img)
    img = enhancer_sharpness.enhance(2.5)
    
    return img

def convert_to_images(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    images = []
    try:
        if ext in ['.png', '.jpg', '.jpeg', '.webp', '.tiff']:
            raw_img = Image.open(uploaded_file).convert('RGB')
            images.append(enhance_image_for_ai(raw_img))
        elif ext == '.pdf':
            pdf = pdfium.PdfDocument(uploaded_file.read())
            page = pdf[0]
            bitmap = page.render(scale=3)
            raw_img = bitmap.to_pil().convert('RGB')
            images.append(enhance_image_for_ai(raw_img))
        elif ext == '.docx':
            doc = Document(uploaded_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            from PIL import ImageDraw
            img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
            ImageDraw.Draw(img).text((20, 20), full_text[:2000], fill=(0, 0, 0))
            images.append(img)
    except Exception as e:
        st.error(f"Error converting file {uploaded_file.name}: {e}")
    return images

def normalize_extracted_values(row_data):
    for key, value in list(row_data.items()):
        if str(value).lower() in ["not found", "null", "none"] or not value: 
            continue
        val_str = str(value).strip()
        if "aadhaar" in key.lower():
            clean_digits = ''.join(filter(str.isdigit, val_str))
            row_data[key] = clean_digits if clean_digits else "Not Found"
        elif "pan" in key.lower() or "identification" in key.lower():
            row_data[key] = re.sub(r'[\s\-]', '', val_str).upper()
        elif "date" in key.lower() or "receipt" in key.lower():
            clean_date = re.sub(r'[\s\.\-\/]', '', val_str)
            if len(clean_date) == 8 and clean_date.isdigit():
                row_data[key] = f"{clean_date[:2]}/{clean_date[2:4]}/{clean_date[4:]}"
    return row_data

def run_regex_validation(extracted_row):
    issues = []
    for k, v in extracted_row.items():
        if str(v).lower() in ["not found", "null", "none"] or not v:
            continue
            
        val_str = re.sub(r'[\s\-]', '', str(v))
        
        # Aadhaar Check
        if "aadhaar" in k.lower():
            if not val_str.isdigit() or len(val_str) != 12:
                issues.append(f"Aadhaar length error ({len(val_str)} digits)")
                
        # PAN Card Check
        elif "pan" in k.lower() or "identification" in k.lower():
            if len(val_str) != 10:
                issues.append(f"PAN length error ({len(val_str)} chars)")
            elif not re.match(r'^[A-Z]{5}[0-9]{4}[A-Z]$', val_str.upper()):
                issues.append("PAN format error")
                
    return "🟢 Validated" if not issues else f"
